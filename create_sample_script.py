"""Create a sample script.docx with interview questions."""

from pathlib import Path

from docx import Document

DATA_DIR = Path(__file__).parent / "data"
SCRIPT_PATH = DATA_DIR / "script.docx"

SAMPLE_QUESTIONS = [
    "What is your full name?",
    "Which city do you currently live in?",
    "What is your highest educational qualification?",
    "How many years of professional experience do you have?",
    "What are your primary technical skills?",
    "Why are you interested in this role?",
    "Describe a challenging project you worked on recently.",
    "What are your salary expectations?",
    "When can you join if selected?",
    "Do you have any questions for us?",
]


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Interview Script", level=1)
    doc.add_paragraph(
        "Aarti will ask only the questions listed below during the voice interview."
    )
    doc.add_paragraph("")
    for i, q in enumerate(SAMPLE_QUESTIONS, start=1):
        doc.add_paragraph(f"{i}. {q}")
    doc.save(SCRIPT_PATH)
    print(f"Created {SCRIPT_PATH} with {len(SAMPLE_QUESTIONS)} questions")


if __name__ == "__main__":
    main()
