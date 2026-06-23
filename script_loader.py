"""Load ordered dialogue steps from script.docx."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

DATA_DIR = Path(__file__).parent / "data"
DEFAULT_SCRIPT_PATH = DATA_DIR / "script.docx"


@dataclass(frozen=True)
class ScriptStep:
    """One line from the script: speak (intro), question (wait for user), or closing."""

    kind: str  # speak | question | closing
    text: str


def _clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[\u200b\ufeff\ufffd\u00ad]", "", text)
    # Drop leading junk before the first letter (handles corrupted docx bytes)
    match = re.search(r"[A-Za-z\u0900-\u097F]", text)
    if match:
        text = text[match.start() :]
    text = re.sub(r"\s+", " ", text).strip()
    text = text.rstrip("\ufffd").strip()
    return text


def _paragraph_text(para) -> str:
    """Join lines inside one docx paragraph into a single spoken line."""
    return _clean_text(para.text.replace("\r\n", " ").replace("\n", " "))


def load_script_steps(path: Path | None = None) -> list[ScriptStep]:
    """Parse script.docx into ordered speak / question / closing steps."""
    script_path = path or DEFAULT_SCRIPT_PATH
    if not script_path.exists():
        raise FileNotFoundError(
            f"Script file not found at {script_path}. "
            "Place your script.docx in backend/data/script.docx"
        )

    doc = Document(script_path)
    paragraphs = [_paragraph_text(p) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = _paragraph_text(para)
                    if text:
                        paragraphs.append(text)

    if not paragraphs:
        raise ValueError(f"No content found in {script_path}")

    steps: list[ScriptStep] = []
    last_idx = len(paragraphs) - 1
    for i, text in enumerate(paragraphs):
        if text.endswith("?"):
            kind = "question"
        elif i == last_idx or re.search(r"dhanyawad", text, re.I):
            kind = "closing"
        else:
            kind = "speak"
        steps.append(ScriptStep(kind=kind, text=text))

    return steps


def load_script_questions(path: Path | None = None) -> list[str]:
    return [s.text for s in load_script_steps(path) if s.kind == "question"]


def load_script_text(path: Path | None = None) -> str:
    return "\n".join(s.text for s in load_script_steps(path))


def script_off_script_reply(steps: list[ScriptStep]) -> str:
    """Pick a scripted line for off-script user questions (no LLM)."""
    for step in reversed(steps):
        if step.kind == "closing":
            return step.text
    for step in steps:
        if step.kind == "speak":
            return step.text
    return "Kripya apna jawab share karein."
